"""Geração do relatório em Word (.docx) via python-docx.

Formato 14×21 cm, fonte da família Palatino ("Palatino Linotype", padrão no
Windows — mesma família da TeX Gyre Pagella usada no PDF). O Word referencia a
fonte pelo nome (não embute), então usamos um nome de fonte amplamente disponível.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from backmedina.report.comum import (
    RESUMO_CABECALHO,
    TITULO,
    legenda_resumo,
    linhas_rastreabilidade,
    linhas_resumo,
)

FONTE = "Palatino Linotype"


def _set_fonte_padrao(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(10.5)


def gerar_docx(
    projeto: dict | None,
    metadados,
    segmentos,
    zi_png: bytes,
    desvio_label: str = "Amostral (n−1)",
    classificacao: dict | None = None,
) -> bytes:
    """Monta o relatório .docx e devolve os bytes."""
    doc = Document()
    _set_fonte_padrao(doc)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(14), Cm(21)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(1.8)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(TITULO)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = FONTE

    nome_proj = (projeto or {}).get("nome", "")
    if nome_proj:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(nome_proj).italic = True

    # --- Rastreabilidade ---
    doc.add_heading("Dados do Projeto (Rastreabilidade)", level=1)
    linhas = linhas_rastreabilidade(projeto, metadados)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    for rotulo, valor in linhas:
        c0, c1 = tbl.add_row().cells
        c0.paragraphs[0].add_run(rotulo).bold = True
        c1.text = valor

    # --- Segmentação: imagem + resumo ---
    doc.add_heading("Segmentação Homogênea", level=1)
    doc.add_picture(io.BytesIO(zi_png), width=Cm(10.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabela = doc.add_table(rows=1, cols=len(RESUMO_CABECALHO))
    tabela.style = "Light Grid Accent 1"
    for i, h in enumerate(RESUMO_CABECALHO):
        tabela.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for linha in linhas_resumo(segmentos):
        cells = tabela.add_row().cells
        for i, v in enumerate(linha):
            cells[i].text = v

    leg = doc.add_paragraph()
    leg.add_run(legenda_resumo(desvio_label)).font.size = Pt(8)

    # --- Classificação estrutural (Horak) ---
    if classificacao and classificacao.get("resumo") is not None:
        resumo = classificacao["resumo"]
        doc.add_heading("Classificação estrutural (Horak)", level=1)
        if classificacao.get("nota"):
            doc.add_paragraph(classificacao["nota"])
        cols = list(resumo.columns)
        tc = doc.add_table(rows=1, cols=len(cols))
        tc.style = "Light Grid Accent 1"
        for i, h in enumerate(cols):
            tc.rows[0].cells[i].paragraphs[0].add_run(str(h)).bold = True
        for _, row in resumo.iterrows():
            cells = tc.add_row().cells
            for i, v in enumerate(row.values):
                cells[i].text = str(v)
        nota = doc.add_paragraph()
        nota.add_run(
            "SCI = camadas superiores · BDI = base · BCI = subleito. Faixas de "
            "Horak em µm (índices ×10). SCI/BDI/BCI escalam com a carga (40 kN)."
        ).font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
