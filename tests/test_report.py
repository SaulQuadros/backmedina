import io
import shutil
from pathlib import Path

import pytest

from backmedina.report.comum import linhas_rastreabilidade, linhas_resumo
from backmedina.report.relatorio_docx import gerar_docx
from backmedina.report.relatorio_pdf import (
    esc,
    gerar_pdf,
    lualatex_disponivel,
    verificar_sem_type1,
)
from backmedina.segmentation.aashto_cumdiff import Segmento


def _png_minimo() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (40, 20), (240, 240, 240)).save(buf, format="PNG")
    return buf.getvalue()


def _segmentos():
    return [
        Segmento(1, 0.0, 160.0, 160.0, 8, 35.6, 3.1, 42.3),
        Segmento(2, 160.0, 480.0, 320.0, 16, 29.5, 2.2, 35.0),
    ]


_PROJ = {
    "nome": "Anel Viário UFJF", "responsavel_tecnico": "Eng. Fulano",
    "empresa_lwd": "SOLOCAP", "tipo_via": "Sistema Arterial Primário",
    "caracteristica_via": "Pista simples", "n_faixas": 2,
    "observacoes": "Teste & validação 100% (com _underscore_)",
}


def test_esc_latex():
    assert esc("a & b_c 100%") == r"a \& b\_c 100\%"


def test_linhas_conteudo():
    linhas = linhas_rastreabilidade(_PROJ, None)
    d = dict(linhas)
    assert d["Projeto"] == "Anel Viário UFJF"
    assert linhas_resumo(_segmentos())[0][0] == "seg_01"


def test_docx_valido():
    from docx import Document

    data = gerar_docx(_PROJ, None, _segmentos(), _png_minimo())
    assert data[:2] == b"PK"
    doc = Document(io.BytesIO(data))
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Retroanálise" in texto
    # tabela de rastreabilidade + resumo presentes
    assert len(doc.tables) >= 2


@pytest.mark.skipif(not lualatex_disponivel(), reason="lualatex ausente")
def test_pdf_sem_type1():
    pdf = gerar_pdf(_PROJ, None, _segmentos(), _png_minimo())
    assert pdf[:4] == b"%PDF"
    ok, tipos = verificar_sem_type1(pdf)
    assert ok, f"PDF contém Type 1: {tipos}"
    # deve usar a família Pagella
    assert any("Pagella" in t for t in tipos)


@pytest.mark.skipif(not lualatex_disponivel(), reason="lualatex ausente")
def test_pdf_zi_sem_type1():
    import numpy as np, pandas as pd
    from backmedina.segmentation.aashto_cumdiff import segmentar, tabela_zi
    from backmedina.report.apresentacao_zi import gerar_pdf_zi
    dist = np.arange(0, 600, 20.0)
    d0 = np.concatenate([np.full(15, 25.0), np.full(15, 60.0)])
    df = pd.DataFrame({"Metros": dist, "D1": d0})
    _, segs = segmentar(df)
    tab, tot = tabela_zi(df)
    pdf = gerar_pdf_zi(tab, tot, segs, _png_minimo())
    assert pdf[:4] == b"%PDF"
    ok, tipos = verificar_sem_type1(pdf)
    assert ok, f"PDF Zi contém Type 1: {tipos}"
